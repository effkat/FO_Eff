import json
import matplotlib.pyplot as plt
import geopandas as gpd
from shapely.geometry import LineString, Point, Polygon, MultiPolygon, box
from fpdf import FPDF
import os
import datetime
import re
import requests
import pandas as pd
import contextily as ctx
import numpy as np
from adjustText import adjust_text
import matplotlib.patheffects as path_effects
from geopy.distance import geodesic
import geopy
from itertools import groupby
from docx import Document
#

# ============================
# 1. GeoJSON-Datei einlesen und analysieren
# ============================

def calculate_length_km(coordinates):
    """Berechnet die Länge eines LineStrings in Kilometern."""
    length_km = 0.0
    for i in range(len(coordinates) - 1):
        point1 = (coordinates[i][1], coordinates[i][0])  # [lat, lon]
        point2 = (coordinates[i + 1][1], coordinates[i + 1][0])  # [lat, lon]
        segment_length = geodesic(point1, point2).km
        length_km += segment_length
    return length_km


def parse_geojson(geojson_file_path):
    """Liest eine GeoJSON-Datei aus und extrahiert relevante Daten."""
    with open(geojson_file_path, "r", encoding="utf-8") as file:
        geojson_data = json.load(file)
    
    features = geojson_data.get("features", [])
    
    geo_data = []
    for feature in features:
        properties = feature.get("properties", {})
        geometry = feature.get("geometry", {})
        line_length_km = None
        if geometry.get("type") == "LineString":
            coordinates = geometry.get("coordinates", [])
            line_length_km = calculate_length_km(coordinates)
        elif geometry.get("type") == "Point":
            coordinates = [geometry.get("coordinates", [])]
        elif geometry.get("type") == "Polygon":
            coordinates = geometry.get("coordinates", [])
        elif geometry.get("type") == "MultiPolygon":
            coordinates = geometry.get("coordinates", [])
        else:
            continue
        
        geo_data.append({
            "name": properties.get("fn:name", "Unnamed"),
            "title": properties.get("fn:title",""),
            "coordinates": coordinates,
            "type": geometry.get("type", ""),
            "length": line_length_km  # Länge in km für LineStrings
        })
    
    return geo_data

# ============================
# 2. Flugroute aus GeoJSON visualisieren
# ============================

def get_street_intersections(polygon_coords):
    # Erstelle das Polygon und seine Außengrenze
    polygon = Polygon(polygon_coords)
    outline = gpd.GeoSeries(polygon.exterior)
    
    # Berechne die Bounding Box des Polygons
    minx, miny, maxx, maxy = polygon.bounds
    
    # Overpass API Anfrage zur Abfrage von Straßen innerhalb der Bounding Box
    overpass_url = "http://overpass-api.de/api/interpreter"
    overpass_query = f"""
    [out:json];
    way["highway"]({miny},{minx},{maxy},{maxx});
    (._;>;);
    out body;
    """
    
    try:
        response = requests.get(overpass_url, params={'data': overpass_query}, timeout=30)
        response.raise_for_status()  # Löst eine Ausnahme bei HTTP-Fehlern aus
        
        if response.text.strip() == "":
            print("Leere Antwort von der Overpass-API erhalten.")
            return None
        
        data = response.json()
    except requests.exceptions.RequestException as e:
        print(f"Fehler beim Abrufen der Daten von Overpass API: {e}")
        return None
    except requests.exceptions.JSONDecodeError:
        print("Fehler beim Dekodieren der JSON-Antwort. Möglicherweise gibt es ein Serverproblem.")
        return None
    
    # Extrahiere Straßen-Geometrien
    elements = data.get("elements", [])
    nodes = {elem["id"]: (elem["lon"], elem["lat"]) for elem in elements if elem["type"] == "node"}
    ways = [
        LineString([nodes[node_id] for node_id in elem.get("nodes", []) if node_id in nodes])
        for elem in elements if elem["type"] == "way" and "nodes" in elem
    ]
    
    if not ways:
        print("Keine Straßen in der Bounding Box gefunden.")
        return None
    
    # Erstelle ein GeoDataFrame
    gdf_streets = gpd.GeoDataFrame(geometry=ways, crs="EPSG:4326")
    
    # Prüfe die Intersections zwischen Straßen und der Außengrenze des Polygons
    intersections = gdf_streets[gdf_streets.geometry.intersects(outline.iloc[0])]
    
    # Konvertiere die Schnittpunkte in Punkte
    intersection_points = []
    for line in intersections.geometry:
        intersection = line.intersection(outline.iloc[0])
        if intersection.geom_type == "Point":
            intersection_points.append(intersection)
        elif intersection.geom_type == "MultiPoint":
            intersection_points.extend(intersection.geoms)
    
    # Erstelle ein GeoDataFrame nur mit Punkten
    gdf_points = gpd.GeoDataFrame(geometry=intersection_points, crs="EPSG:4326")
    
    
    return gdf_points


def extract_points_from_geo_data(geo_data):
    """Erstellt einen DataFrame mit Punktgeometrien aus der geo_data-Liste, 
       wobei nur Punkte mit einem gültigen 'title' berücksichtigt werden."""
    point_data = []

    for item in geo_data:
        if item["type"] == "Point" and item["title"]:  # Nur Punkte mit vorhandenem Titel aufnehmen
            coordinates = item["coordinates"]
            if len(coordinates) == 1:  # Punktkoordinaten als Liste enthalten
                lon, lat = round(coordinates[0][0], 5), round(coordinates[0][1], 5)
                point_data.append({
                    "Point": item["title"],
                    "Lat": lat,
                    "Lon": lon
                })

    df = pd.DataFrame(point_data, columns=["Point", "Lat", "Lon"])
    return df

def calculate_extreme_points_dataframe(geo_data):
    """Ermittelt die nördlichsten, östlichsten, südlichsten und westlichsten Punkte entlang der LineStrings."""
    extreme_points = {
        "northernmost": None,
        "easternmost": None,
        "southernmost": None,
        "westernmost": None
    }
    
    for item in geo_data:
        if item["type"] == "LineString":
            coordinates = item.get("coordinates", [])
            for point in coordinates:
                lat, lon = point[1], point[0]
                
                if extreme_points["northernmost"] is None or lat > extreme_points["northernmost"][1]:
                    extreme_points["northernmost"] = (lon, lat)
                if extreme_points["easternmost"] is None or lon > extreme_points["easternmost"][0]:
                    extreme_points["easternmost"] = (lon, lat)
                if extreme_points["southernmost"] is None or lat < extreme_points["southernmost"][1]:
                    extreme_points["southernmost"] = (lon, lat)
                if extreme_points["westernmost"] is None or lon < extreme_points["westernmost"][0]:
                    extreme_points["westernmost"] = (lon, lat)
    
    data = [
        {"Point": "Northernmost", "Lat": round(extreme_points["northernmost"][1], 5), "Lon": round(extreme_points["northernmost"][0], 5)},
        {"Point": "Easternmost", "Lat": round(extreme_points["easternmost"][1], 5), "Lon": round(extreme_points["easternmost"][0], 5)},
        {"Point": "Southernmost", "Lat": round(extreme_points["southernmost"][1], 5), "Lon": round(extreme_points["southernmost"][0], 5)},
        {"Point": "Westernmost", "Lat": round(extreme_points["westernmost"][1], 5), "Lon": round(extreme_points["westernmost"][0], 5)}
    ]
    return pd.DataFrame(data)

def find_nearest_PIS(geo_data, PIS_geojson_path):
    """
    Findet den nächstgelegenen Punkt aus der Punkte-GeoJSON-Datei zum LineString aus geo_data.
    
    :param geo_data: Liste mit GeoJSON-Feature-Daten (bereits eingelesen und verarbeitet)
    :param points_geojson_path: Pfad zur GeoJSON-Datei mit Punkten
    :return: Der nächstgelegene Punkt mit Code, Name, Ort, Koordinaten und Distanz in km
    """
    # Lade die Punkte-GeoJSON-Datei
    with open(PIS_geojson_path, "r", encoding="utf-8") as file:
        points_data = json.load(file)
    
    points = [(feature["properties"].get("Code", "Unknown"),
               feature["properties"].get("Name des Krankenhauses", "Unnamed"),
               feature["properties"].get("Ort", "Unknown"),
               feature["geometry"]["coordinates"]) 
              for feature in points_data.get("features", [])
              if feature.get("geometry", {}).get("type") == "Point"]
    
    # Extrahiere den Linestring aus geo_data
    line = next((feature for feature in geo_data if feature["type"] == "LineString"), None)
    
    if not line:
        raise ValueError("Kein LineString in geo_data gefunden.")
    
    line_coords = line["coordinates"]
    
    # Finde den nächstgelegenen Punkt
    min_distance = float("inf")
    nearest_PIS = None
    nearest_info = {}
    
    for code, name, location, point in points:
        for segment_start, segment_end in zip(line_coords[:-1], line_coords[1:]):
            dist = geopy.distance.distance((point[1], point[0]), (segment_start[1], segment_start[0])).km
            if dist < min_distance:
                min_distance = dist
                nearest_PIS = point
                PIS_info = {
                    "code": code,
                    "name": name,
                    "location": location,
                    "coordinates": nearest_PIS,
                    "distance_km": min_distance
                }
    
    return PIS_info

def buffer_and_get_bbox(point, buffer_distance=50):
    """Buffert einen Punkt und gibt die Bounding Box zurück."""
    buffered_point = Point(point).buffer(buffer_distance)
    minx, miny, maxx, maxy = buffered_point.bounds
    return minx, miny, maxx, maxy

def calculate_width_height(bbox):
    """Berechnet width und height basierend auf der BBOX-Größe."""
    minx, miny, maxx, maxy = bbox
    width = int((maxx - minx) * 10000)  # Skalierung anpassen
    height = int((maxy - miny) * 10000)  # Skalierung anpassen
    return max(1, width), max(1, height)

def get_wms_info(bbox, width, height, layer):
    """
    Führt die GetFeatureInfo-Abfrage für einen WMS-Layer durch.
    Die Antwort wird als Klartext zurückgegeben.
    """
    url = "https://uas-betrieb.de/geoservices/dipul/wms"
    params = {
        "service": "WMS",
        "version": "1.3.0",
        "request": "GetFeatureInfo",
        "query_layers": f"dipul:{layer}",
        "layers": f"dipul:{layer}",
        "bbox": f"{bbox[1]},{bbox[0]},{bbox[3]},{bbox[2]}",  # Reihenfolge für EPSG:4326 (lat, lon)
        "width": width,
        "height": height,
        "srs": "EPSG:4326",
        "feature_count": 10,
        "i": width // 2,
        "j": height // 2
    }
    response = requests.get(url, params=params)
    return response.text

def extract_name(response_text):
    """
    Durchsucht den WMS-Response-Text zeilenweise und gibt den Text hinter "name =" zurück.
    """
    for line in response_text.splitlines():
        line = line.strip()
        if line.startswith("name ="):
            # Teilt die Zeile an "=" und gibt den rechtsstehenden (getrimmten) Teil zurück.
            return line.split("=", 1)[1].strip()
    return None

def get_highway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("H")]
    
    highway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["bundesautobahnen", "bundesstrassen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)
            
            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue
            
            # Setze den Alias gemäß des Layers
            if layer == "bundesstrassen":
                alias = "Federal Highway (Bundesstraße)"
            elif layer == "bundesautobahnen":
                alias = "State Highway (Autobahn)"
            else:
                alias = ""
            
            highway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "feature_name": feature_name,
                "alias": alias
            })
    return highway_info

def get_railway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("R")]
    
    railway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["bahnanlagen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            
            # Setze den Alias gemäß des Layers
            if layer == "bahnanlagen":
                alias = "railway line"
            else:
                alias = ""
            
            railway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "alias": alias
            })
    return railway_info

def get_waterway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("W")]
    
    waterway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["binnenwasserstrassen","seewasserstrassen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)

            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue
            
            # Setze den Alias gemäß des Layers
            if layer == "binnenwasserstrassen":
                alias = "federal waterway"
            elif layer == "seewasserstrassen":
                alias = "seawaterway"
            else:
                alias = ""
            
            waterway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "feature_name": feature_name,
                "alias": alias
            })
    return waterway_info


def get_nature_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "N" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("N")]
    
    nature_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["nationalparks", "naturschutzgebiete", "vogelschutzgebiete"]:
            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)
            
            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue
            
            # Setze den Alias gemäß des Layers
            if layer == "nationalparks":
                alias = "national park"
            elif layer == "naturschutzgebiete":
                alias = "conservation area"
            elif layer == "vogelschutzgebiete":
                alias = "bird sanctuary"
            else:
                alias = ""
            
            nature_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "feature_name": feature_name,
                "alias": alias
            })
    return nature_info




def generate_flight_map(geo_data, image_path):
    """Erstellt eine quadratische Karte der Flugroute mit Voyager-Basemap und Koordinatenrahmen."""
    
    line_strings = []
    points = []
    polygons = []
    labels = []
    polygon_labels = []
    
    for entry in geo_data:
        coords = entry["coordinates"]
        geom_type = entry["type"]  # Standard: LineString
        label = entry.get("title", "")
        name = entry.get("name", "") 

        if geom_type == "Polygon":
            polygon = Polygon(coords[0])  # Außenring des Polygons extrahieren
            polygons.append(polygon)
            if name:
                polygon_labels.append((polygon.centroid.x, polygon.centroid.y, name))
        elif geom_type == "MultiPolygon":
            for poly_coords in coords:
                polygon = Polygon(poly_coords[0])  # Alle Polygone aus MultiPolygon extrahieren
                polygons.append(polygon)
                if name:
                    polygon_labels.append((polygon.centroid.x, polygon.centroid.y, name))
        elif len(coords) > 1:
            line_strings.append(LineString(coords))  # Mehrere Punkte -> LineString
        elif label:  # Punkt wird nur gespeichert, wenn label nicht leer ist
            lon, lat = coords[0][:2]  # Einzelner Punkt
            points.append(Point(lon, lat))
            labels.append((lon, lat, label))  # Label speichern

    # GeoDataFrames erstellen
    gdf_lines = gpd.GeoDataFrame(geometry=line_strings, crs="EPSG:4326")
    gdf_points = gpd.GeoDataFrame(geometry=points, crs="EPSG:4326")
    gdf_polygons = gpd.GeoDataFrame(geometry=polygons, crs="EPSG:4326")

    # Basemap vorbereiten (Web-Mapping: EPSG 3857)
    gdf_lines_3857 = gdf_lines.to_crs(epsg=3857)
    gdf_points_3857 = gdf_points.to_crs(epsg=3857)
    gdf_polygons_3857 = gdf_polygons.to_crs(epsg=3857)

    # A4 Hochformat mit 3/4 Platz für die Karte
    fig = plt.figure(figsize=(8.3, 11.7))  # A4 Hochformat
    ax = fig.add_axes([0.05, 0.05, 0.7, 0.7])  # Platzierung unten (3/4 der Seite)

    # Linien zuerst zeichnen
    gdf_lines_3857.plot(ax=ax, color="blue", linewidth=1.5, label="Flight Route", zorder=2)  

    # Polygone im Hintergrund zeichnen
    gdf_polygons_3857.plot(ax=ax, facecolor="yellow", alpha=0.5, edgecolor="gold", linewidth=0.7, label="DDAs", zorder=1)  

    # Punkte als Marker über die Linien zeichnen
    x_coords = gdf_points_3857.geometry.x
    y_coords = gdf_points_3857.geometry.y
    ax.scatter(x_coords, y_coords, facecolors="orangered", s=100, marker='.', edgecolors="black", linewidth=0.8, label="Waypoints", zorder=3)

    # Labels für Punkte hinzufügen
    texts = []
    for (lon, lat, label) in labels:
        if label:
            point_3857 = gpd.GeoSeries([Point(lon, lat)], crs="EPSG:4326").to_crs(epsg=3857).geometry[0]
            text = ax.text(
                point_3857.x, point_3857.y, label,
                fontsize=8, fontweight="bold", ha='center', va='bottom',
                color='black',
                path_effects=[
                    path_effects.Stroke(linewidth=3, foreground="white"),  # Weißer Rand
                    path_effects.Normal()
                ]
            )
            texts.append(text)

    # Labels für Polygone hinzufügen
    for (lon, lat, name) in polygon_labels:
        point_3857 = gpd.GeoSeries([Point(lon, lat)], crs="EPSG:4326").to_crs(epsg=3857).geometry[0]
        text = ax.text(
            point_3857.x, point_3857.y, name,
            fontsize=8, fontweight="bold", ha='center', va='bottom',
            color='black',
            path_effects=[
                path_effects.Stroke(linewidth=3, foreground="white"),  # Weißer Rand
                path_effects.Normal()
            ]
        )
        texts.append(text)

    # Anpassung der Pfeile mit weißem Hintergrund
    adjust_text(
        texts, ax=ax, only_move={'points': 'y', 'text': 'y'},
        arrowprops=dict(arrowstyle='-', color='black', lw=0.8, alpha=1, 
                        path_effects=[path_effects.withStroke(linewidth=2, foreground="white")])
    )

    # Voyager-Basemap hinzufügen
    ctx.add_basemap(ax, source=ctx.providers.CartoDB.Voyager, zoom=12, attribution=False)
    
    # Achsenbeschriftung in WGS84 Koordinaten (EPSG:4326) mit weniger Ticks
    xticks_3857 = np.linspace(ax.get_xlim()[0]*1.001, ax.get_xlim()[1], num=5)
    yticks_3857 = np.linspace(ax.get_ylim()[0], ax.get_ylim()[1], num=5)
    xticks_wgs84 = [gpd.GeoSeries([Point(x, yticks_3857[0])], crs="EPSG:3857").to_crs("EPSG:4326").geometry[0].x for x in xticks_3857]
    yticks_wgs84 = [gpd.GeoSeries([Point(xticks_3857[0], y)], crs="EPSG:3857").to_crs("EPSG:4326").geometry[0].y for y in yticks_3857]
    ax.set_xticks(xticks_3857)
    ax.set_yticks(yticks_3857)
    ax.set_xticklabels([f"{tick:.2f}" for tick in xticks_wgs84], rotation=0)
    ax.set_yticklabels([f"{tick:.2f}" for tick in yticks_wgs84], rotation=90)

    # Karte speichern
    App_Name = geo_data[0]["name"] if geo_data else "Unknown Route"
    suffix=extract_suffix(App_Name)
    plt.title(f"Figure A{suffix}.2", loc="right", fontdict={"fontsize":9, "verticalalignment":'bottom'})
    fig.savefig(image_path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def generate_TO_map(geo_data, image_path2):
    """Erstellt eine Karte mit einem einzelnen Punkt und Buffern, zentriert in einem 400m x 200m Ausschnitt."""
    
    points = []
    labels = []
    
    for entry in geo_data:
        coords = entry["coordinates"]
        label = entry.get("title", "")  # Nur der Titel des Punkts wird verwendet
        
        if label:  
            lon, lat = coords[0][:2]  # Einzelner Punkt
            points.append(Point(lon, lat))
            labels.append((lon, lat, label))  # Titel speichern
    
    if not points:
        raise ValueError("Keine gültigen Punkte in den Daten gefunden.")

    # GeoDataFrame erstellen
    gdf_points = gpd.GeoDataFrame(geometry=points, crs="EPSG:4326")
    gdf_points_3857 = gdf_points.to_crs(crs=3857)

    # Buffer-Zonen erstellen (100m = rot, 150m = gelb)
    buffer_100m = gdf_points_3857.buffer(50)
    buffer_150m = gdf_points_3857.buffer(75)
    buffer_int = buffer_150m.to_crs(crs=4326)

    polygon_coords = list(buffer_int.geometry.iloc[0].exterior.coords)
    intersections = get_street_intersections(polygon_coords)
    intersections = intersections.to_crs(crs=3857)
    
    # Kartenfigur erstellen
    fig, ax = plt.subplots(figsize=(6, 3))

    # Buffers plotten
    buffer_150m.plot(ax=ax, facecolor="none", edgecolor="yellow", linewidth=0.5, label="150m")
    buffer_100m.plot(ax=ax, facecolor="none", edgecolor="red", linewidth=0.5, label="100m")
    intersections.plot(ax=ax, facecolor="whitesmoke", markersize=50,marker='^',edgecolors="red", linewidth=0.8, zorder=1)


    # Punkt plotten
    gdf_points_3857.plot(ax=ax, facecolors="blue", markersize=60, marker=u"$\mathbb{\otimes}$", edgecolors="blue", label="TO-Site", linewidth=0.1, zorder=3)

    # Achseneinstellungen - Ausschnitt 400m x 200m
    x_center, y_center = gdf_points_3857.geometry.x[0], gdf_points_3857.geometry.y[0]
    ax.set_xlim([x_center - 300, x_center + 300])  # 400m breit
    ax.set_ylim([y_center - 200, y_center + 200])  # 200m hoch

    # Satellitenkarte als Hintergrund
    ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, attribution=False,zoom=18)

     # Achsen ausblenden
    ax.set_xticks([])
    ax.set_yticks([])
    ax.set_frame_on(False)
    # Karte speichern
    App_Name = geo_data[0]["name"] if geo_data else "Unknown Route"
    suffix=extract_suffix(App_Name)
    plt.title(f"Figure A{suffix}.2", loc="right", fontdict={"fontsize":5, "verticalalignment":'bottom'})
    fig.savefig(image_path2, dpi=300, bbox_inches="tight")
    plt.close(fig)

def extract_suffix(route_name):
    """
    Extrahiert die letzten ein oder zwei Ziffern aus dem String,
    sofern sie zwischen 1 und 9 liegen.
    """
    match = re.search(r"([1-9]{1,2})$", route_name)  # Suche nach 1 oder 2 Ziffern am Ende (1-9)
    if match:
        return match.group(1)  # Gibt die gefundene Ziffer oder Ziffern zurück
    return ""  # Falls keine passende Ziffer gefunden wird, leere Zeichenkette zurückgeben

def get_bundesland_from_geo_data(geo_data):
    """Nimmt geo_data als Input und gibt das Bundesland der ersten Koordinate aus."""
    for entry in geo_data:
        if entry["coordinates"]:
            lon, lat, _ = entry["coordinates"][0]  # GeoJSON speichert Koordinaten als [Längengrad, Breitengrad, Höhe]
            
            url = f"https://overpass-api.de/api/interpreter?data=[out:json];is_in({lat},{lon});area._[\"admin_level\"=\"4\"];out;"
            try:
                response = requests.get(url, timeout=5)
                response.raise_for_status()
                data = response.json()
                if "elements" in data and data["elements"]:
                    return data["elements"][0].get("tags", {}).get("name", "Unbekannt")
            except requests.exceptions.RequestException as e:
                print(f"Fehler bei der Anfrage: {e}")
            except ValueError:
                print("Fehler: Ungültige JSON-Antwort von der API erhalten.")
    
    return "Keine gültigen Koordinaten gefunden"


# ============================
# 3. PDF-Bericht generieren
# ============================

class FlightReportPDF(FPDF):
    """Definiert das PDF-Layout für den Flugbericht."""
    def __init__(self, route_name="Unknown Route", logo_path=None):
        super().__init__()
        self.route_name = route_name  # Speichert den Namen des LineStrings
        self.logo_path = logo_path
        self.logo_width = 45
        self.ignore_header = False
        self.set_margins(20, 15, 20)
        self.set_auto_page_break(auto=True, margin=15)
        self.add_font('Montserrat', '', r'Montserrat\static\Montserrat-Regular.ttf', uni=True)
        self.add_font('Montserrat', 'B', r'Montserrat\static\Montserrat-Bold.ttf', uni=True)
        self.add_font('Montserrat', 'I', r'Montserrat\static\Montserrat-Italic.ttf', uni=True)

    def header(self):
        """Setzt den Header mit dem Namen der Route."""
        if not self.ignore_header:
            if self.logo_path and os.path.exists(self.logo_path):
                self.image(self.logo_path, x=150, y=8, w=self.logo_width)
            self.set_y(15)
            self.set_font("Montserrat", "I", 12)
            self.set_text_color(50, 50, 50)
            geo_data = parse_geojson(geojson_file_path)
            route_name = geo_data[0]["name"] if geo_data else "Unknown Route"
            self.cell(0, 10, f"Beagle Systems FO-Document: {route_name}", ln=True, align="L")
            self.set_draw_color(20, 20, 20)
            self.line(20, 25, 190, 25)
            self.ln(5)

    def footer(self):
        """Setzt die Fußzeile mit einer Linie und der Seitenzahl."""
        self.set_y(-15)  # Position für die Fußzeile
        self.set_draw_color(20, 20, 20)  # Farbe der Linie (schwarz)
        # Linie oberhalb der Fußzeile zeichnen
        self.line(10, self.get_y() - 2, 190, self.get_y() - 2)  # Linie über der Fußzeile

        self.set_font("Montserrat", "I", 8)  # Schriftart setzen
        self.cell(0, 10, f"Page {self.page_no()}", align="C")  # Fußzeilentext setzen


    def chapter_title(self, title):
        self.set_font("Montserrat", "B", 12)
        self.cell(0, 8, title, ln=True, align="L")
        self.ln(5)
    
    def chapter_title2(self, title):
        self.set_font("Montserrat", "", 12)
        self.cell(0, 5, title, ln=True, align="L")
        self.ln(5)

    def chapter_body(self, text):
        self.set_font("Montserrat", "", 10)
        self.multi_cell(0, 5, text)
        self.ln()
    
    def add_table(self, data, col_widths):
        """Fügt eine Tabelle in das PDF ein."""
        self.set_font("Montserrat", "B", 9)
        headers = ["Route distances", "2D-Distance", "3D-Distance", "Estimated Flight Time"]
        for i, header in enumerate(headers):
            self.cell(col_widths[i], 10, header, border=1, align='C')
        self.ln()
        
        self.set_font("Montserrat", "", 9)
        for row in data:
            for i, item in enumerate(row):
                if i == 3:  # Die Spalte mit den Zeilenumbrüchen
                    self.multi_cell(col_widths[i], 6, str(item), border=1, align='C')
                else:
                    self.cell(col_widths[i], 18, str(item), border=1, align='C')


    def coordinate_table(self, df, col_widths, ignore_header=False):
        """Fügt eine Koordinaten-Tabelle in das PDF ein."""
        self.set_font("Montserrat", "B", 11)
        headers = ["Point", "Lat", "Lon"]

        # Kopfzeile nur erstellen, wenn ignore_header False ist
        if not ignore_header:
            for i, header in enumerate(headers):
                self.cell(col_widths[i], 10, header, border=1, align='C')
            self.ln()

        self.set_font("Montserrat", "", 11)
    
        # Datenzeilen aus dem DataFrame einfügen
        for _, row in df.iterrows():
            self.cell(col_widths[0], 10, str(row["Point"]), border=1, align='C')
            self.cell(col_widths[1], 10, f"{row['Lat']:.5f}", border=1, align='C')
            self.cell(col_widths[2], 10, f"{row['Lon']:.5f}", border=1, align='C')
            self.ln()

    def coordinate_table2(self, df, col_widths, ignore_header=False):
        """Fügt eine Koordinaten-Tabelle in das PDF ein."""
        self.set_font("Montserrat", "B", 11)
        headers = ["Point", "Lat", "Lon"]

        # Kopfzeile nur erstellen, wenn ignore_header False ist
        if not ignore_header:
            for i, header in enumerate(headers):
                self.cell(col_widths[i], 10, header, border=1, align='C')
            self.ln()

        self.set_font("Montserrat", "", 11)

        # Sortierung: "TO" zuerst, dann alphabetisch
        df_sorted = df.copy()
        df_sorted["is_TO"] = df_sorted["Point"].apply(lambda x: "TO" in x)
        df_sorted = df_sorted.sort_values(by=["is_TO", "Point"], ascending=[False, True]).drop(columns=["is_TO"])

        # Datenzeilen aus dem DataFrame einfügen
        for _, row in df_sorted.iterrows():
            self.cell(col_widths[0], 10, str(row["Point"]), border=1, align='C')
            self.cell(col_widths[1], 10, f"{row['Lat']:.5f}", border=1, align='C')
            self.cell(col_widths[2], 10, f"{row['Lon']:.5f}", border=1, align='C')
            self.ln()

        

    def chapter_body2(self, text):
        self.set_font("Montserrat", "B", 11)
        self.multi_cell(0, 7, text)
        self.ln()
    
    def chapter_body3(self, text):
        self.set_font("Montserrat", "", 10)
        self.multi_cell(0, 3, text)
        self.ln()

    def copyright_text(self, text):
        self.set_font("Montserrat","I",11)
        self.multi_cell(0,5,text,align='C')
        self.ln()


# ============================
# 4. Funktion zum Erstellen des PDF-Berichts
# ============================

def generate_pdf_report(geojson_file_path, output_pdf_path):
    """Erstellt den PDF-Bericht aus der GeoJSON-Datei."""
    geo_data = parse_geojson(geojson_file_path)
    route_name = geo_data[0]["name"] if geo_data else "Unknown Route"
    route_length= geo_data[0]["length"] if geo_data else "No length"
    latitude, longitude = None, None

    for entry in geo_data:
        if entry.get("type") == "Point" and "title" in entry and "TO" in entry["title"]:
            if "coordinates" in entry and isinstance(entry["coordinates"], list) and entry["coordinates"]:
                lon, lat = entry["coordinates"][0][:2]  # Extrahiere die ersten beiden Werte (Lon, Lat)
                latitude = f"{round(lat, 5)}°N"
                longitude = f"{round(lon, 5)}°E"
            break  # Stoppt die Schleife, sobald das erste passende Feature gefunden wurde

    # Bild der Flugroute erstellen
    image_path = "flight_route.png"
    generate_flight_map(geo_data, image_path)

    image_path2="TO-Site.png"
    generate_TO_map(geo_data,image_path2)

    # PDF erstellen
    pdf = FlightReportPDF(route_name=route_name, logo_path="Beagle.png")
    pdf.add_page()

    suffix = extract_suffix(route_name)

    df = pd.read_csv("3DFaktor.csv")  # CSV-Datei laden
    bundesland = get_bundesland_from_geo_data(geo_data)  # Bundesland bestimmen

    D_Faktor = df[df.iloc[:, 1] == bundesland].iloc[0, 2] if bundesland in df.iloc[:, 1].values else "Nicht gefunden"

    # Berechnungen für Tabelle
    line_length_km = route_length # Beispielwert, muss berechnet werden
    equivalent_3d_distance = round(line_length_km * D_Faktor, 2)
    flight_time_mk1 = round(equivalent_3d_distance*1000 / 24 /60,0)
    flight_time_mk2 = round(equivalent_3d_distance*1000 / 28/60,0)
    flight_time_octo = round(equivalent_3d_distance*1000 / 3/60,0)
    
    table_data = [["From take-off/landing site and back", 
               f"{line_length_km:.2f} km", 
               f"{equivalent_3d_distance:.2f} km", 
               f"\n"
               f"{int(flight_time_octo)} min\n"
               f"\n"]]
    col_widths = [60, 35, 35, 45]

    pdf.chapter_body2(f"Appendix {suffix} to the Flight Operation Document {route_name[:10]} \n" 
                     f"{datetime.date.today().strftime('%d.%m.%Y')}\n"
                     f"ISSUE 00\n"
                     f"Flight Route {suffix}")

    pdf.chapter_title("1. Route Overview")
    pdf.chapter_body(f"Figure A{suffix}.1 gives a general overview of the mission.")
    pdf.image(image_path, x=15, y=85, w=180)

    pdf.add_page()
    pdf.chapter_title("2. Flight Distances and Times")
    pdf.add_table(table_data, col_widths)
    pdf.chapter_body("")
    pdf.chapter_title("3. Takeoff / Landing Site")
    pdf.chapter_body(f"Takeoff and Landing occur at coordinates ({latitude} {longitude}), see Figure A{suffix}.2 for details.")
    pdf.image(image_path2,x=20, y=107, w=170)

    pdf.add_page()
    pdf.chapter_title("4. Detailed Information")

    pdf.chapter_title2("Detailed population density assessment")
    pdf.chapter_body("None required.")
    pdf.chapter_body("")

    pdf.chapter_title2("Individual Approvals")
    pdf.chapter_body("None required.")
    pdf.chapter_body("")

    pdf.chapter_title2("Highway")
    highways = get_highway_names(geo_data)
    if highways == []:
        pdf.chapter_body("None required")
    else:
        pdf.chapter_body(f"See Figure A{suffix}.1 and under 'Coordinates'")
        highways.sort(key=lambda x: x["point_title"])
        for H in highways:
            pdf.chapter_body3(f"At point {H['point_title']} the {H['alias']} {H['feature_name']} is crossed at an angle of 90°.")
        pdf.chapter_body("")

    pdf.chapter_title2("Railway Lines")
    railways = get_railway_names(geo_data)
    if railways == []:
        pdf.chapter_body("None required.")
    else:
        pdf.chapter_body(f"See Figure A{suffix}.1 and under 'Coordinates'")
        railways.sort(key=lambda x: x["point_title"])
        for R in railways:
            pdf.chapter_body3(f"At point {R['point_title']} a {R['alias']} is crossed at an angle of 90°.")
        pdf.chapter_body("")

    pdf.chapter_title2("Federal Waterways")
    waterways = get_waterway_names(geo_data)
    if waterways == []:
        pdf.chapter_body("None required.")
    else:
        pdf.chapter_body(f"See Figure A{suffix}.1 and under 'Coordinates'")
        waterways.sort(key=lambda x: x["point_title"])
        for i in range(0, len(waterways) - 1, 2):
            point_from = waterways[i]
            point_to = waterways[i + 1]
            pdf.chapter_body3(
                f"From point {point_from['point_title']} to point {point_to['point_title']} the route conducts over the {point_from['alias']} {point_from['feature_name']}."
        )

    pdf.chapter_body("")

    pdf.chapter_title2("Power Lines")
    pdf.chapter_body("None required.")
    pdf.chapter_body("")

    pdf.chapter_title2("Closest public interest site (PIS)")
    PISinfo = find_nearest_PIS(geo_data, 'PIS.geojson')
    pdf.chapter_body3(f"{PISinfo['name']}, PIS {PISinfo['code']}:  {PISinfo['distance_km']:.2f} km")
    pdf.chapter_body("")

    pdf.chapter_title2("Natural Preserves")
    natures = get_nature_names(geo_data)
    if natures == []:
        pdf.chapter_body("None required.")
    else:
        pdf.chapter_body(f"See Figure A{suffix}.1 and under 'Coordinates'")

        # Sortiere zuerst global nach alias, dann nach name und schließlich nach point_title
        natures.sort(key=lambda x: (x["alias"], x["feature_name"], x["point_title"]))

            # Gruppiere die Einträge nach (alias, name)
        for (alias, name), group in groupby(natures, key=lambda x: (x["alias"], x["feature_name"])):
            # Erzeuge eine Liste aus der aktuellen Gruppe
            group_list = list(group)
            # Optional: Ausgabe einer Überschrift für die Gruppe
            pdf.chapter_body3(f"{alias} - ({name}):")

            # Falls gewünscht: Sortiere die Gruppe nochmals explizit nach point_title (ist meist schon durch den globalen Sortiervorgang gegeben)
            group_list.sort(key=lambda x: x["point_title"])

            # Erstelle Zweier-Paare: von element 0 zu 1, von element 2 zu 3, etc.
            for i in range(0, len(group_list) - 1, 2):
                point_from = group_list[i]
                point_to = group_list[i + 1]
                pdf.chapter_body3(f"            -   From point {point_from['point_title']} to point {point_to['point_title']}.")
        
        # Leere Zeile als Trenner zwischen Gruppen (optional)
    pdf.chapter_body("")


    pdf.add_page()
    pdf.chapter_title("5. Coordinates")
    col_widths = [60, 50, 50]  
    bbox_df = calculate_extreme_points_dataframe(geo_data)
    pdf.coordinate_table(bbox_df, col_widths,ignore_header=False)
    cord_table = extract_points_from_geo_data(geo_data)
    pdf.coordinate_table2(cord_table, col_widths,ignore_header=True)

    pdf.add_page()
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.chapter_title("")
    pdf.copyright_text("Information stared in this paper is proprietary and belongs exclusivley to Beagle Systems and may not be reproduced without any prior notice the knowledge or a written consent approved by Beagle Systems.")
    pdf.copyright_text("www.beaglesystems.com")

    pdf.output(output_pdf_path)
    print(f"PDF report generated: {output_pdf_path}")


def generate_docx_report(geojson_file_path, output_docx_path):
    document = Document()
    
    

# ============================
# 5. Skript ausführen
# ============================

if __name__ == "__main__":
    geojson_file_path = "24-39188_20250203.geojson"
    output_pdf_path = "flight_report.pdf"
    
    generate_pdf_report(geojson_file_path, output_pdf_path)