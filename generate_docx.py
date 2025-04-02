from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
import os
from fpdf import FPDF
from fpdf import HTMLMixin
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import pandas as pd
import json
import matplotlib.pyplot as plt
import geopandas as gpd
from shapely.geometry import LineString, Point, Polygon, MultiPolygon, box
from fpdf import FPDF
import os
import datetime
import re
import requests
import contextily as ctx
import numpy as np
from adjustText import adjust_text
import matplotlib.patheffects as path_effects
from geopy.distance import geodesic
import geopy
from itertools import groupby




# ============================
# 1. GeoJSON-Datei einlesen und analysieren
# ============================

def calculate_length_km(coordinates):
    """Berechnet die Länge eines LineStrings in Kilometern."""
    length_km = 0.0
    for i in range(len(coordinates) - 1):
        point1 = (coordinates[i][1], coordinates[i][0])  # [lat, lon]
        point2 = (coordinates[i + 1][1], coordinates[i + 1][0])  # [lat, lon]
        segment_length = geodesic(point1, point2).km
        length_km += segment_length
    return length_km


def parse_geojson(geojson_file_path):
    """Liest eine GeoJSON-Datei aus und extrahiert relevante Daten."""
    with open(geojson_file_path, "r", encoding="utf-8") as file:
        geojson_data = json.load(file)
    
    features = geojson_data.get("features", [])
    
    geo_data = []
    for feature in features:
        properties = feature.get("properties", {})
        geometry = feature.get("geometry", {})
        line_length_km = None
        if geometry.get("type") == "LineString":
            coordinates = geometry.get("coordinates", [])
            line_length_km = calculate_length_km(coordinates)
        elif geometry.get("type") == "Point":
            coordinates = [geometry.get("coordinates", [])]
        elif geometry.get("type") == "Polygon":
            coordinates = geometry.get("coordinates", [])
        elif geometry.get("type") == "MultiPolygon":
            coordinates = geometry.get("coordinates", [])
        else:
            continue
        
        geo_data.append({
            "name": properties.get("fn:name", "Unnamed"),
            "title": properties.get("fn:title",""),
            "coordinates": coordinates,
            "type": geometry.get("type", ""),
            "length": line_length_km  # Länge in km für LineStrings
        })
    
    return geo_data

# ============================
# 2. Flugroute aus GeoJSON visualisieren
# ============================

def get_street_intersections(polygon_coords):
    # Erstelle das Polygon und seine Außengrenze
    polygon = Polygon(polygon_coords)
    outline = gpd.GeoSeries(polygon.exterior)
    
    # Berechne die Bounding Box des Polygons
    minx, miny, maxx, maxy = polygon.bounds
    
    # Overpass API Anfrage zur Abfrage von Straßen innerhalb der Bounding Box
    overpass_url = "http://overpass-api.de/api/interpreter"
    overpass_query = f"""
    [out:json];
    way["highway"]({miny},{minx},{maxy},{maxx});
    (._;>;);
    out body;
    """
    
    try:
        response = requests.get(overpass_url, params={'data': overpass_query}, timeout=30)
        response.raise_for_status()  # Löst eine Ausnahme bei HTTP-Fehlern aus
        
        if response.text.strip() == "":
            print("Leere Antwort von der Overpass-API erhalten.")
            return None
        
        data = response.json()
    except requests.exceptions.RequestException as e:
        print(f"Fehler beim Abrufen der Daten von Overpass API: {e}")
        return None
    except requests.exceptions.JSONDecodeError:
        print("Fehler beim Dekodieren der JSON-Antwort. Möglicherweise gibt es ein Serverproblem.")
        return None
    
    # Extrahiere Straßen-Geometrien
    elements = data.get("elements", [])
    nodes = {elem["id"]: (elem["lon"], elem["lat"]) for elem in elements if elem["type"] == "node"}
    ways = [
        LineString([nodes[node_id] for node_id in elem.get("nodes", []) if node_id in nodes])
        for elem in elements if elem["type"] == "way" and "nodes" in elem
    ]
    
    if not ways:
        print("Keine Straßen in der Bounding Box gefunden.")
        return None
    
    # Erstelle ein GeoDataFrame
    gdf_streets = gpd.GeoDataFrame(geometry=ways, crs="EPSG:4326")
    
    # Prüfe die Intersections zwischen Straßen und der Außengrenze des Polygons
    intersections = gdf_streets[gdf_streets.geometry.intersects(outline.iloc[0])]
    
    # Konvertiere die Schnittpunkte in Punkte
    intersection_points = []
    for line in intersections.geometry:
        intersection = line.intersection(outline.iloc[0])
        if intersection.geom_type == "Point":
            intersection_points.append(intersection)
        elif intersection.geom_type == "MultiPoint":
            intersection_points.extend(intersection.geoms)
    
    # Erstelle ein GeoDataFrame nur mit Punkten
    gdf_points = gpd.GeoDataFrame(geometry=intersection_points, crs="EPSG:4326")
    
    
    return gdf_points

def extract_points_from_geo_data(geo_data):
    """Erstellt einen DataFrame mit Punktgeometrien aus der geo_data-Liste, 
       wobei nur Punkte mit einem gültigen 'title' berücksichtigt werden."""
    point_data = []

    for item in geo_data:
        if item["type"] == "Point" and item["title"]:  # Nur Punkte mit vorhandenem Titel aufnehmen
            coordinates = item["coordinates"]
            if len(coordinates) == 1:  # Punktkoordinaten als Liste enthalten
                lon, lat = round(coordinates[0][0], 5), round(coordinates[0][1], 5)
                point_data.append({
                    "Point": item["title"],
                    "Lat": lat,
                    "Lon": lon
                })

    df = pd.DataFrame(point_data, columns=["Point", "Lat", "Lon"])
    return df

def calculate_extreme_points_dataframe(geo_data):
    """Ermittelt die nördlichsten, östlichsten, südlichsten und westlichsten Punkte entlang der LineStrings."""
    extreme_points = {
        "northernmost": None,
        "easternmost": None,
        "southernmost": None,
        "westernmost": None
    }
    
    for item in geo_data:
        if item["type"] == "LineString":
            coordinates = item.get("coordinates", [])
            for point in coordinates:
                lat, lon = point[1], point[0]
                
                if extreme_points["northernmost"] is None or lat > extreme_points["northernmost"][1]:
                    extreme_points["northernmost"] = (lon, lat)
                if extreme_points["easternmost"] is None or lon > extreme_points["easternmost"][0]:
                    extreme_points["easternmost"] = (lon, lat)
                if extreme_points["southernmost"] is None or lat < extreme_points["southernmost"][1]:
                    extreme_points["southernmost"] = (lon, lat)
                if extreme_points["westernmost"] is None or lon < extreme_points["westernmost"][0]:
                    extreme_points["westernmost"] = (lon, lat)
    
    data = [
        {"Point": "Northernmost", "Lat": round(extreme_points["northernmost"][1], 5), "Lon": round(extreme_points["northernmost"][0], 5)},
        {"Point": "Easternmost", "Lat": round(extreme_points["easternmost"][1], 5), "Lon": round(extreme_points["easternmost"][0], 5)},
        {"Point": "Southernmost", "Lat": round(extreme_points["southernmost"][1], 5), "Lon": round(extreme_points["southernmost"][0], 5)},
        {"Point": "Westernmost", "Lat": round(extreme_points["westernmost"][1], 5), "Lon": round(extreme_points["westernmost"][0], 5)}
    ]
    return pd.DataFrame(data)

def find_nearest_PIS(geo_data, PIS_geojson_path):
    """
    Findet den nächstgelegenen Punkt aus der Punkte-GeoJSON-Datei zum LineString aus geo_data.
    
    :param geo_data: Liste mit GeoJSON-Feature-Daten (bereits eingelesen und verarbeitet)
    :param points_geojson_path: Pfad zur GeoJSON-Datei mit Punkten
    :return: Der nächstgelegene Punkt mit Code, Name, Ort, Koordinaten und Distanz in km
    """
    # Lade die Punkte-GeoJSON-Datei
    with open(PIS_geojson_path, "r", encoding="utf-8") as file:
        points_data = json.load(file)
    
    points = [(feature["properties"].get("Code", "Unknown"),
               feature["properties"].get("Name des Krankenhauses", "Unnamed"),
               feature["properties"].get("Ort", "Unknown"),
               feature["geometry"]["coordinates"]) 
              for feature in points_data.get("features", [])
              if feature.get("geometry", {}).get("type") == "Point"]
    
    # Extrahiere den Linestring aus geo_data
    line = next((feature for feature in geo_data if feature["type"] == "LineString"), None)
    
    if not line:
        raise ValueError("Kein LineString in geo_data gefunden.")
    
    line_coords = line["coordinates"]
    
    # Finde den nächstgelegenen Punkt
    min_distance = float("inf")
    nearest_PIS = None
    nearest_info = {}
    
    for code, name, location, point in points:
        for segment_start, segment_end in zip(line_coords[:-1], line_coords[1:]):
            dist = geopy.distance.distance((point[1], point[0]), (segment_start[1], segment_start[0])).km
            if dist < min_distance:
                min_distance = dist
                nearest_PIS = point
                PIS_info = {
                    "code": code,
                    "name": name,
                    "location": location,
                    "coordinates": nearest_PIS,
                    "distance_km": min_distance
                }
    
    return PIS_info

def buffer_and_get_bbox(point, buffer_distance=50):
    """Buffert einen Punkt und gibt die Bounding Box zurück."""
    buffered_point = Point(point).buffer(buffer_distance)
    minx, miny, maxx, maxy = buffered_point.bounds
    return minx, miny, maxx, maxy

def calculate_width_height(bbox):
    """Berechnet width und height basierend auf der BBOX-Größe."""
    minx, miny, maxx, maxy = bbox
    width = int((maxx - minx) * 10000)  # Skalierung anpassen
    height = int((maxy - miny) * 10000)  # Skalierung anpassen
    return max(1, width), max(1, height)

def get_wms_info(bbox, width, height, layer):
    """
    Führt die GetFeatureInfo-Abfrage für einen WMS-Layer durch.
    Die Antwort wird als Klartext zurückgegeben.
    """
    url = "https://uas-betrieb.de/geoservices/dipul/wms"
    params = {
        "service": "WMS",
        "version": "1.3.0",
        "request": "GetFeatureInfo",
        "query_layers": f"dipul:{layer}",
        "layers": f"dipul:{layer}",
        "bbox": f"{bbox[1]},{bbox[0]},{bbox[3]},{bbox[2]}",  # Reihenfolge für EPSG:4326 (lat, lon)
        "width": width,
        "height": height,
        "srs": "EPSG:4326",
        "feature_count": 10,
        "i": width // 2,
        "j": height // 2
    }
    response = requests.get(url, params=params)
    return response.text

def extract_name(response_text):
    """
    Durchsucht den WMS-Response-Text zeilenweise und gibt den Text hinter "name =" zurück.
    """
    for line in response_text.splitlines():
        line = line.strip()
        if line.startswith("name ="):
            # Teilt die Zeile an "=" und gibt den rechtsstehenden (getrimmten) Teil zurück.
            return line.split("=", 1)[1].strip()
    return None

def get_highway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("H")]
    
    highway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["bundesautobahnen", "bundesstrassen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)
            
            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue
            
            # Setze den Alias gemäß des Layers
            if layer == "bundesstrassen":
                alias = "Federal Highway (Bundesstraße)"
            elif layer == "bundesautobahnen":
                alias = "State Highway (Autobahn)"
            else:
                alias = ""
            
            highway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "feature_name": feature_name,
                "alias": alias
            })
    return highway_info

def get_railway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("R")]
    
    railway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["bahnanlagen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            
            # Setze den Alias gemäß des Layers
            if layer == "bahnanlagen":
                alias = "railway line"
            else:
                alias = ""
            
            railway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "alias": alias
            })
    return railway_info

def get_waterway_names(geo_data):
    # Filter: nur Punkte, deren Titel mit "H" beginnt.
    filtered_points = [point for point in geo_data if point.get("title", "").startswith("W")]
    
    waterway_info = []
    for point in filtered_points:
        bbox = buffer_and_get_bbox(point["coordinates"])
        width, height = calculate_width_height(bbox)
        
        for layer in ["binnenwasserstrassen","seewasserstrassen"]:
            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)

            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue
            
            # Setze den Alias gemäß des Layers
            if layer == "binnenwasserstrassen":
                alias = "federal waterway"
            elif layer == "seewasserstrassen":
                alias = "seawaterway"
            else:
                alias = ""
            
            waterway_info.append({
                "point_title": point.get("title"),
                "layer": layer,
                "feature_name": feature_name,
                "alias": alias
            })
    return waterway_info

def get_nature_names(geo_data):

    nature_info = []
    for layer in ["nationalparks", "naturschutzgebiete", "vogelschutzgebiete"]:
        # Filter: nur Punkte, deren Titel mit "N" beginnt und mit layer endet.
        filtered_points = [point for point in geo_data if point.get("title", "").startswith("N") and point.get("title", "").endswith(layer)]

        for point in filtered_points:
            bbox = buffer_and_get_bbox(point["coordinates"])
            width, height = calculate_width_height(bbox)

            response_text = get_wms_info(bbox, width, height, layer)
            feature_name = extract_name(response_text)

            # Wenn kein feature_name extrahiert wurde, überspringe diesen Fall
            if feature_name is None:
                continue

            # Setze den Alias gemäß des Layers
            if layer == "nationalparks":
                alias = "national park"
            elif layer == "naturschutzgebiete":
                alias = "conservation area"
            elif layer == "vogelschutzgebiete":
                alias = "bird sanctuary"
            else:
                alias = ""

            nature_info.append({
            "point_title": point.get("title")[:2],
            "layer": layer,
            "feature_name": feature_name,
            "alias": alias
            })
                
           
    return nature_info

def generate_flight_map(geo_data, image_path):
    """Erstellt eine quadratische Karte der Flugroute mit Voyager-Basemap und Koordinatenrahmen."""
    
    line_strings = []
    points = []
    polygons = []
    labels = []
    polygon_labels = []
    
    for entry in geo_data:
        coords = entry["coordinates"]
        geom_type = entry["type"]  # Standard: LineString
        label = entry.get("title", "")
        name = entry.get("name", "") 

        if geom_type == "Polygon":
            polygon = Polygon(coords[0])  # Außenring des Polygons extrahieren
            polygons.append(polygon)
            if name:
                polygon_labels.append((polygon.centroid.x, polygon.centroid.y, name))
        elif geom_type == "MultiPolygon":
            for poly_coords in coords:
                polygon = Polygon(poly_coords[0])  # Alle Polygone aus MultiPolygon extrahieren
                polygons.append(polygon)
                if name:
                    polygon_labels.append((polygon.centroid.x, polygon.centroid.y, name))
        elif len(coords) > 1:
            line_strings.append(LineString(coords))  # Mehrere Punkte -> LineString
        elif label:  # Punkt wird nur gespeichert, wenn label nicht leer ist
            lon, lat = coords[0][:2]  # Einzelner Punkt
            points.append(Point(lon, lat))
            labels.append((lon, lat, label))  # Label speichern

    # GeoDataFrames erstellen
    gdf_lines = gpd.GeoDataFrame(geometry=line_strings, crs="EPSG:4326")
    gdf_points = gpd.GeoDataFrame(geometry=points, crs="EPSG:4326")
    gdf_polygons = gpd.GeoDataFrame(geometry=polygons, crs="EPSG:4326")

    # Basemap vorbereiten (Web-Mapping: EPSG 3857)
    gdf_lines_3857 = gdf_lines.to_crs(epsg=3857)
    gdf_points_3857 = gdf_points.to_crs(epsg=3857)
    gdf_polygons_3857 = gdf_polygons.to_crs(epsg=3857)

    # A4 Hochformat mit 3/4 Platz für die Karte
    fig = plt.figure(figsize=(8.3, 11.7))  # A4 Hochformat
    ax = fig.add_axes([0.05, 0.05, 0.7, 0.7])  # Platzierung unten (3/4 der Seite)

    # Linien zuerst zeichnen
    gdf_lines_3857.plot(ax=ax, color="blue", linewidth=1.5, label="Flight Route", zorder=2)  

    # Polygone im Hintergrund zeichnen
    gdf_polygons_3857.plot(ax=ax, facecolor="yellow", alpha=0.5, edgecolor="gold", linewidth=0.7, label="DDAs", zorder=1)  

    # Punkte als Marker über die Linien zeichnen
    x_coords = gdf_points_3857.geometry.x
    y_coords = gdf_points_3857.geometry.y
    ax.scatter(x_coords, y_coords, facecolors="orangered", s=100, marker='.', edgecolors="black", linewidth=0.8, label="Waypoints", zorder=3)

    # Labels für Punkte hinzufügen
    texts = []
    for (lon, lat, label) in labels:
        if label:
            point_3857 = gpd.GeoSeries([Point(lon, lat)], crs="EPSG:4326").to_crs(epsg=3857).geometry[0]
            text = ax.text(
                point_3857.x, point_3857.y, label,
                fontsize=8, fontweight="bold", ha='center', va='bottom',
                color='black',
                path_effects=[
                    path_effects.Stroke(linewidth=3, foreground="white"),  # Weißer Rand
                    path_effects.Normal()
                ]
            )
            texts.append(text)

    # Labels für Polygone hinzufügen
    for (lon, lat, name) in polygon_labels:
        point_3857 = gpd.GeoSeries([Point(lon, lat)], crs="EPSG:4326").to_crs(epsg=3857).geometry[0]
        text = ax.text(
            point_3857.x, point_3857.y, name,
            fontsize=8, fontweight="bold", ha='center', va='bottom',
            color='black',
            path_effects=[
                path_effects.Stroke(linewidth=3, foreground="white"),  # Weißer Rand
                path_effects.Normal()
            ]
        )
        texts.append(text)

    # Anpassung der Pfeile mit weißem Hintergrund
    adjust_text(
        texts, ax=ax, only_move={'points': 'y', 'text': 'y'},
        arrowprops=dict(arrowstyle='-', color='black', lw=0.8, alpha=1, 
                        path_effects=[path_effects.withStroke(linewidth=2, foreground="white")])
    )

    # Voyager-Basemap hinzufügen
    ctx.add_basemap(ax, source=ctx.providers.CartoDB.Voyager, zoom=12, attribution=False)
    
    # Achsenbeschriftung in WGS84 Koordinaten (EPSG:4326) mit weniger Ticks
    xticks_3857 = np.linspace(ax.get_xlim()[0]*1.001, ax.get_xlim()[1], num=5)
    yticks_3857 = np.linspace(ax.get_ylim()[0], ax.get_ylim()[1], num=5)
    xticks_wgs84 = [gpd.GeoSeries([Point(x, yticks_3857[0])], crs="EPSG:3857").to_crs("EPSG:4326").geometry[0].x for x in xticks_3857]
    yticks_wgs84 = [gpd.GeoSeries([Point(xticks_3857[0], y)], crs="EPSG:3857").to_crs("EPSG:4326").geometry[0].y for y in yticks_3857]
    ax.set_xticks(xticks_3857)
    ax.set_yticks(yticks_3857)
    ax.set_xticklabels([f"{tick:.2f}" for tick in xticks_wgs84], rotation=0)
    ax.set_yticklabels([f"{tick:.2f}" for tick in yticks_wgs84], rotation=90)

    # Karte speichern
    App_Name = geo_data[0]["name"] if geo_data else "Unknown Route"
    suffix=extract_suffix(App_Name)
    plt.title(f"Figure A{suffix}.2", loc="right", fontdict={"fontsize":9, "verticalalignment":'bottom'})
    fig.savefig(image_path, dpi=300, bbox_inches="tight")
    plt.close(fig)

def generate_TO_map(geo_data, image_path2):
    """Erstellt eine Karte mit einem einzelnen Punkt und Buffern, zentriert in einem 400m x 200m Ausschnitt."""
    
    points = []
    labels = []
    
    for entry in geo_data:
        coords = entry["coordinates"]
        label = entry.get("title", "")  # Nur der Titel des Punkts wird verwendet
        
        if label:  
            lon, lat = coords[0][:2]  # Einzelner Punkt
            points.append(Point(lon, lat))
            labels.append((lon, lat, label))  # Titel speichern
    
    if not points:
        raise ValueError("Keine gültigen Punkte in den Daten gefunden.")

    # GeoDataFrame erstellen
    gdf_points = gpd.GeoDataFrame(geometry=points, crs="EPSG:4326")
    gdf_points_3857 = gdf_points.to_crs(crs=3857)

    # Buffer-Zonen erstellen (100m = rot, 150m = gelb)
    buffer_100m = gdf_points_3857.buffer(50)
    buffer_150m = gdf_points_3857.buffer(75)
    buffer_int = buffer_150m.to_crs(crs=4326)

    polygon_coords = list(buffer_int.geometry.iloc[0].exterior.coords)
    intersections = get_street_intersections(polygon_coords)
    if intersections is not None :
        intersections = intersections.to_crs(crs=3857)
    
    # Kartenfigur erstellen
    fig, ax = plt.subplots(figsize=(6, 3))

    # Buffers plotten
    buffer_150m.plot(ax=ax, facecolor="none", edgecolor="yellow", linewidth=0.5, label="150m")
    buffer_100m.plot(ax=ax, facecolor="none", edgecolor="red", linewidth=0.5, label="100m")
    
    if intersections is not None:
        intersections.plot(ax=ax, facecolor="whitesmoke", markersize=50,marker='^',edgecolors="red", linewidth=0.8, zorder=1)


    # Punkt plotten
    gdf_points_3857.plot(ax=ax, facecolors="blue", markersize=60, marker=u"$\mathbb{\otimes}$", edgecolors="blue", label="TO-Site", linewidth=0.1, zorder=3)

    # Achseneinstellungen - Ausschnitt 400m x 200m
    x_center, y_center = gdf_points_3857.geometry.x[0], gdf_points_3857.geometry.y[0]
    ax.set_xlim([x_center - 300, x_center + 300])  # 400m breit
    ax.set_ylim([y_center - 200, y_center + 200])  # 200m hoch

    # Satellitenkarte als Hintergrund
    ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, attribution=False,zoom=18)

     # Achsen ausblenden
    ax.set_xticks([])
    ax.set_yticks([])
    ax.set_frame_on(False)
    # Karte speichern
    App_Name = geo_data[0]["name"] if geo_data else "Unknown Route"
    suffix=extract_suffix(App_Name)
    plt.title(f"Figure A{suffix}.2", loc="right", fontdict={"fontsize":5, "verticalalignment":'bottom'})
    fig.savefig(image_path2, dpi=300, bbox_inches="tight")
    plt.close(fig)

def extract_suffix(route_name):
    """
    Extrahiert die letzten ein oder zwei Ziffern aus dem String,
    sofern sie zwischen 1 und 9 liegen.
    """
    match = re.search(r"([1-9]{1,2})$", route_name)  # Suche nach 1 oder 2 Ziffern am Ende (1-9)
    if match:
        return match.group(1)  # Gibt die gefundene Ziffer oder Ziffern zurück
    return ""  # Falls keine passende Ziffer gefunden wird, leere Zeichenkette zurückgeben

def get_bundesland_from_geo_data(geo_data):
    """Nimmt geo_data als Input und gibt das Bundesland der ersten Koordinate aus."""
    for entry in geo_data:
        if entry["coordinates"]:
            lon, lat, _ = entry["coordinates"][0]  # GeoJSON speichert Koordinaten als [Längengrad, Breitengrad, Höhe]
            
            url = f"https://overpass-api.de/api/interpreter?data=[out:json];is_in({lat},{lon});area._[\"admin_level\"=\"4\"];out;"
            try:
                response = requests.get(url, timeout=5)
                response.raise_for_status()
                data = response.json()
                if "elements" in data and data["elements"]:
                    return data["elements"][0].get("tags", {}).get("name", "Unbekannt")
            except requests.exceptions.RequestException as e:
                print(f"Fehler bei der Anfrage: {e}")
            except ValueError:
                print("Fehler: Ungültige JSON-Antwort von der API erhalten.")
    
    return "Keine gültigen Koordinaten gefunden"



##### # DOCX-Report erstellen

# muss ich hieraus ne class machen?

def generate_docx_report(output_path):

    def header(route_name):
        """Setzt den Header mit dem Namen der Route."""
        #add header

        section = document.sections[0]
        header = section.header

        table = header.add_table(rows=1, cols=2, width= Inches(8)) 
        table.columns[0].width = Inches(5)  
        table.columns[1].width = Inches(3)  

        # Get the first row of the table
        row = table.rows[0]

        # Add text to the left cell
        left_cell = row.cells[0]
        left_cell.text = 'FO-Document: ' + route_name

        # Add image to the right cell
        right_cell = row.cells[1]

        # Insert an image into the right cell
        right_cell.paragraphs[0].clear()  
        run = right_cell.paragraphs[0].add_run()
        run.add_picture('beagle.png', width=Inches(1.5))  

        # Align the text to the left for the left cell
        left_cell.paragraphs[0].alignment = 0  

        # Align the image to the right side of the cell
        right_cell.paragraphs[0].alignment = 2  

        # Add a line below the header
        line_paragraph = header.add_paragraph()
        line_paragraph.paragraph_format.space_before = 0
        
        insertHR(line_paragraph)
    
    #generate page numbers

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)


    def add_page_number(run):
        run.text = 'Page '
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def insertHR(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:top')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def generate_distance_table(bundesland, route_length):
        df = pd.read_csv("3DFaktor.csv")  # CSV-Datei laden
        bundesland = get_bundesland_from_geo_data(geo_data)  # Bundesland bestimmen
        D_Faktor = df[df.iloc[:, 1] == bundesland].iloc[0, 2] if bundesland in df.iloc[:, 1].values else "Nicht gefunden"

        # Berechnungen für Tabelle
        line_length_km = route_length # Beispielwert, muss berechnet werden
        equivalent_3d_distance = round(line_length_km * D_Faktor, 2)
        flight_time_mk1 = round(equivalent_3d_distance*1000 / 24 /60,0)
        flight_time_mk2 = round(equivalent_3d_distance*1000 / 28/60,0)
        flight_time_octo = round(equivalent_3d_distance*1000 / 3/60,0)
        
        headers = ["Route distances", "2D-Distance", "3D-Distance", "Estimated Flight Time"]
        table_data = ["From take-off/landing site and back", 
                   f"{line_length_km:.2f} km", 
                   f"{equivalent_3d_distance:.2f} km", 
                   #f"\n"
                   f"{int(flight_time_octo)} min\n"
                   #f"\n"
                   ]

        table = document.add_table(rows=2, cols=4)
        heading_cells = table.rows[0].cells
        data_cells = table.rows[1].cells

        for i in range(len(headers)):
            heading_cells[i].paragraphs[0].add_run(headers[i]).bold = True
            heading_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(len(table_data)):
            data_cells[i].paragraphs[0].add_run(table_data[i])
            data_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.style = 'Table Grid'

    def generate_coordinate_table(geo_data):

        bbox_df = calculate_extreme_points_dataframe(geo_data)
        cord_table = extract_points_from_geo_data(geo_data)

        table = document.add_table(bbox_df.shape[0]+1+cord_table.shape[0], bbox_df.shape[1])

        j_0 = 0
        # add the header rows
        for j in range(bbox_df.shape[-1]):
            table.cell(0,j).text = bbox_df.columns[j]

        # add the bbox_df rows
        for i in range(bbox_df.shape[0]):
            for j in range(bbox_df.shape[-1]):
                table.cell(i+1,j).text = str(bbox_df.values[i,j])
                j_0+=1
        
        #add the cord_table rows
        for i in range(cord_table.shape[0]):
            for j in range(cord_table.shape[-1]):
                table.cell(i+1,j+j_0).text = str(cord_table.values[i,j])

        table.style = 'Table Grid'

        
    def footer():

        #add line
        footer = section.footer
        footer_line_para = footer.paragraphs[0]
        insertHR(footer_line_para)

        #add page number
        footer_para = footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
        add_page_number(footer_para.add_run())

    # variables

    geojson_file_path = "24-39188_20250203.geojson" 
    geo_data = parse_geojson(geojson_file_path)
    route_name = geo_data[0]["name"] if geo_data else "Unknown Route"
    route_length= geo_data[0]["length"] if geo_data else "No length"
    latitude, longitude = None, None

    for entry in geo_data:
        if entry.get("type") == "Point" and "title" in entry and "TO" in entry["title"]:
            if "coordinates" in entry and isinstance(entry["coordinates"], list) and entry["coordinates"]:
                lon, lat = entry["coordinates"][0][:2]  # Extrahiere die ersten beiden Werte (Lon, Lat)
                latitude = f"{round(lat, 5)}°N"
                longitude = f"{round(lon, 5)}°E"
            break  # Stoppt die Schleife, sobald das erste passende Feature gefunden wurde

    # Bild der Flugroute erstellen
    image_path = "flight_route.png"
    generate_flight_map(geo_data, image_path)

    image_path2="TO-Site.png"
    generate_TO_map(geo_data,image_path2)

    suffix = extract_suffix(route_name)
    issue = "00"
    bundesland = get_bundesland_from_geo_data(geo_data)


    
    # docx laden

    document = Document() # für Tabellen Style template.docx erstellen
    sections = document.sections

    for section in sections: 
        # set margins
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)

    # Set the default font for all paragraphs
    style = document.styles['Normal']
    font = style.font
    font.name = 'Montserrat'           
    font.size = Pt(12)            
    font.color.rgb = RGBColor(0,0,0)
    
    # Set the default font for all paragraphs
    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Montserrat'           
    font.size = Pt(14)            
    font.color.rgb = RGBColor(0,0,0)
    font.bold = True

    # Set the default font for all paragraphs
    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Montserrat'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0,0,0)
    font.bold = False

    

    
    

    # demo variables
    
    # route_name = "Unknown Route"
    # suffix = "28"
    # 
    # image_path_overview = "flight_route.png"
    # bundesland = "Berlin"
    # route_length = 1000
    # image_path2="TO-Site.png"
    # latitude = '54.05822'
    # longitude = '13.86399'
    
    header(route_name)

    #intro paragraph

    intro_para = document.add_paragraph()
    intro_para_text = f"Appendix {suffix} to the Flight Operation Document {route_name[:10]} \n"+ f"{datetime.date.today().strftime('%d.%m.%Y')}\n"+f"ISSUE {issue}\n"+f"Flight Route {suffix}"
    intro_para.add_run(intro_para_text)

    #route overview

    document.add_heading('1. Route Overview', 1)

    route_overview_para = document.add_paragraph()
    route_overview_para.add_run(f"Figure A{suffix}.1 gives a general overview of the mission.\n")
    image1_run = route_overview_para.add_run()
    image1_run.add_picture(image_path, width=Inches(19.5))
    

    #Flight distances and Times

    document.add_heading('2. Flight Distances and Times', 1)

    #flight_distances_para = document.add_paragraph()

    generate_distance_table(bundesland, route_length)

    # Takeoff and Landing Sites
    document.add_heading('3. Takeoff and Landing Sites', 1)
    TOL_para = document.add_paragraph()
    TOL_para.add_run(f"Takeoff and Landing occur at coordinates ({latitude} {longitude}), see Figure A{suffix}.2 for details.")
    image2_run = TOL_para.add_run()
    image2_run.add_picture(image_path2, width=Inches(19.5))
    
    document.add_heading('4. Detailed Information', 1)


    # Population Density
    document.add_heading('Detailed Population Density Assessment',2)

    population_para = document.add_paragraph()
    population_para_text = f"None required"
    population_para.add_run(population_para_text)

    # Individual Approvals

    document.add_heading('Individual Approvals',2)
    individual_approvals_para = document.add_paragraph()
    individual_approvals_para_text = f"None required"
    individual_approvals_para.add_run(individual_approvals_para_text)

    # Highways

    document.add_heading('Highway',2)
    highway_para = document.add_paragraph()
    highway_para_text = f"None required"
    highway_para.add_run(highway_para_text)

    # Railways    

    document.add_heading('Railway Lines',2)
    railway_para = document.add_paragraph()

    railways = get_railway_names(geo_data)
    if railways == []:
        railway_para_text = f"None required"
        railway_para.add_run(railway_para_text)
    else:
        railway_para.add_run(f"See Figure A{suffix}.1 and under 'Coordinates'. \n")
        railways.sort(key=lambda x: x["point_title"])
        for R in railways:
            railway_para.add_run(f"At point {R['point_title']} a {R['alias']} is crossed at an angle of 90°. \n")
       
    

    # Waterways
    document.add_heading('Federal Waterways',2)
    waterways_para = document.add_paragraph()
    waterways = get_waterway_names(geo_data)

    if waterways == []:
        waterways_para.add_run("None required.")
    else:
        waterways_para.add_run(f"See Figure A{suffix}.1 and under 'Coordinates' \n")
        waterways.sort(key=lambda x: x["point_title"])
        for i in range(0, len(waterways) - 1, 2):
            point_from = waterways[i]
            point_to = waterways[i + 1]
            waterways_para.add_run(
                f"From point {point_from['point_title']} to point {point_to['point_title']} the route conducts over the {point_from['alias']} {point_from['feature_name']}. \n"
        )

    # Power Lines

    document.add_heading('Power Lines',2)
    power_lines_para = document.add_paragraph()
    power_lines_para_text = f"None required"
    power_lines_para.add_run(power_lines_para_text)

    #PIS

    document.add_heading('Closest public interest site (PIS)',2)
    pis_para = document.add_paragraph()
    PISinfo = find_nearest_PIS(geo_data, 'PIS.geojson')
    pis_para.add_run(f"{PISinfo['name']}, PIS {PISinfo['code']}:  {PISinfo['distance_km']:.2f} km")


    # Nature Preserves
    document.add_heading('Natural Preserves',2)
    nature_para = document.add_paragraph()

    natures = get_nature_names(geo_data)
    if natures == []:
        nature_para.add_run("None required.")
    else:
        nature_para.add_run(f"See Figure A{suffix}.1 and under 'Coordinates'")

        # Sortiere zuerst global nach alias, dann nach name und schließlich nach point_title
        natures.sort(key=lambda x: (x["alias"], x["feature_name"], x["point_title"]))

        # Gruppiere die Einträge nach (alias, name)
        for (alias, name), group in groupby(natures, key=lambda x: (x["alias"], x["feature_name"])):
            # Erzeuge eine Liste aus der aktuellen Gruppe
            group_list = list(group)
            print(group_list)
            # Optional: Ausgabe einer Überschrift für die Gruppe
            nature_para.add_run(f"{alias} - ({name}):")

            # Falls gewünscht: Sortiere die Gruppe nochmals explizit nach point_title (ist meist schon durch den globalen Sortiervorgang gegeben)
            group_list.sort(key=lambda x: x["point_title"])

            # Erstelle Zweier-Paare: von element 0 zu 1, von element 2 zu 3, etc.
            for i in range(0, len(group_list) - 1, 2):
                point_from = group_list[i]
                point_to = group_list[i + 1] 
                nature_para.add_run(f"            -   From point {point_from['point_title']} to point {point_to['point_title']}.")


    # Coordinates   
    document.add_heading('5. Coordinates', 1)

    generate_coordinate_table(geo_data)


    footer()

    document.save(output_path)

    print(f"Report saved as {output_docx_path}")

output_docx_path = "flight_report.docx"

generate_docx_report(output_docx_path)